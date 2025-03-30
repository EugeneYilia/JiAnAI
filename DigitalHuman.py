import logging
import os
import shutil
import subprocess
import warnings
import zipfile
from datetime import datetime

import gradio as gr
import jieba
import requests
import torch
import whisper
from PIL import Image as PILImage  # 用于获取图像分辨率
from TTS.api import TTS
from opencc import OpenCC

from utils.CutVoice import trim_tail_by_energy_and_gradient

# --------------------------------------------------------------------------
# 忽略部分警告
warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", category=FutureWarning)
jieba.setLogLevel(jieba.logging.WARN)

# --------------------------------------------------------------------------
# 初始化必要的文件夹
UPLOADS_DIR = "uploads"
os.makedirs(UPLOADS_DIR, exist_ok=True)

RECOGNIZED_DIR = "recognized"
os.makedirs(RECOGNIZED_DIR, exist_ok=True)
RECOGNIZED_EXPORT_DIR = "recognized_export"
os.makedirs(RECOGNIZED_EXPORT_DIR, exist_ok=True)

# --------------------------------------------------------------------------
# 日志器设置 - 在脚本/主入口处执行
app_logger = logging.getLogger("app_logger")
app_logger.setLevel(logging.INFO)

while app_logger.handlers:
    app_logger.handlers.pop()

fh = logging.FileHandler(os.path.join(UPLOADS_DIR, "upload.log"), encoding="utf-8")
fh.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(message)s"))
app_logger.addHandler(fh)

app_logger.propagate = False

# 新增语音识别日志器，记录每一次识别结果
asr_logger = logging.getLogger("asr_logger")
asr_logger.setLevel(logging.INFO)
while asr_logger.handlers:
    asr_logger.handlers.pop()
asr_fh = logging.FileHandler(os.path.join(UPLOADS_DIR, "recognized.log"), encoding="utf-8")
asr_fh.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(message)s"))
asr_logger.addHandler(asr_fh)
asr_logger.propagate = False

def filter_connection_reset_error(record: logging.LogRecord) -> bool:
    msg = record.getMessage()
    if "ConnectionResetError" in msg or "forcibly closed by the remote host" in msg:
        return False
    return True

app_logger.addFilter(filter_connection_reset_error)

asyncio_logger = logging.getLogger("asyncio")
asyncio_logger.propagate = False
asyncio_logger.addFilter(filter_connection_reset_error)

# --------------------------------------------------------------------------
# 注册 RAdam 类及初始化 OpenCC
cc = OpenCC('t2s')

# --------------------------------------------------------------------------
# CSS 样式设置
material_css = """
@import url('https://fonts.googleapis.com/css?family=Roboto:400,500,700&display=swap');

html {
  overflow-y: scroll; /* 始终显示滚动条 */
}

:root {
  --md-primary: #1976d2;
  --md-primary-dark: #1565c0;
  --md-secondary: #424242;
  --md-text: #212121;
  --md-text-on-primary: #ffffff;
  --md-border-radius: 8px;
  --md-transition: 0.3s ease;
}

/* 背景部分 */
html, body, .gradio-container {
  margin: 0;
  padding: 0;
  font-family: 'Roboto', sans-serif;
  color: var(--md-text);
  background-color: #f6e2b3 !important;
  background-image:
    repeating-linear-gradient(45deg, rgba(0,0,0,0.03), rgba(0,0,0,0.03) 1px, transparent 1px, transparent 8px),
    linear-gradient(rgba(255,255,255,0.35), rgba(255,255,255,0.35)),
    url("https://raw.githubusercontent.com/EugeneYilia/JiAnAI/master/assets/images/freemasonry.png");
  background-size: auto, cover, cover;
  background-repeat: repeat, no-repeat, no-repeat;
  background-position: center, center, center;
  background-attachment: fixed, fixed, fixed;
}

/* 主要内容容器背景 */
.tabs, .tabitem, .gr-box, .gr-group, .gr-row, .gr-column {
  background-color: rgba(255, 255, 255, 0.7) !important;
  border-radius: var(--md-border-radius) !important;
  box-shadow: 0 2px 8px rgba(0,0,0,0.08);
  margin-top: 8px !important;
  padding: 12px !important;
}

/* 输入区域、文件上传、音频组件 */
.gr-textbox, .gr-file, .gr-audio {
  background-color: #ffffff !important;
  border-radius: var(--md-border-radius) !important;
  box-shadow: 0 1px 3px rgba(0,0,0,0.1);
  padding: 8px !important;
}
.gr-textbox textarea {
  min-height: 100px !important;
  background-color: #fff !important;
  border: 1px solid #ccc !important;
  border-radius: var(--md-border-radius) !important;
}

/* 自定义按钮样式 */
.gr-button:hover {
  background-color: var(--md-primary-dark) !important;
  box-shadow: 0 4px 8px rgba(0,0,0,0.2);
}
.gr-button:active {
  transform: scale(0.98);
}

/* Tab 按钮选中状态及点击响应 */
.tabs button.selected {
  color: var(--md-primary) !important;
  border-bottom: 3px solid var(--md-primary) !important;
  background-color: transparent !important;
}
.tabs button:active {
  transform: scale(0.98);
}

/* 🚩 强制底部 Footer 所有内容为爱马仕橙 */
footer button.show-api,
footer button.show-api *,
footer button.settings,
footer button.settings *,
footer a.built-with,
footer a.built-with *,
footer .built-with * {
  color: #FF7F00 !important;
  fill: #FF7F00 !important;
  opacity: 1 !important;
  filter: none !important;
  text-shadow: none !important;
}

/* 🎯 专门强制修改 Settings 的齿轮图标 */
footer button.settings svg,
footer button.settings path {
  fill: #FF7F00 !important;
  stroke: #FF7F00 !important;
  color: #FF7F00 !important;
}
"""

# --------------------------------------------------------------------------
# 全局设置及初始化 TTS/ASR
def safe_register_all_globals():
    torch.serialization._allowed_globals = {
        "__builtin__": set(dir(__builtins__)),
        "builtins": set(dir(__builtins__)),
    }
    torch.serialization._allowed_globals.update({
        "TTS.utils.radam": {"RAdam"},
        "TTS.vocoder.utils.generic": {"ADAM"},
        "TTS.models.tacotron.loss": {"TacotronLoss"},
        "TTS.vocoder.models.wavernn": {"Wavernn"},
    })

safe_register_all_globals()

MODEL_NAME = "tts_models/zh-CN/baker/tacotron2-DDC-GST"
tts = TTS(model_name=MODEL_NAME, progress_bar=True, gpu=False)

# 删除原来的固定加载 asr 模型，改为动态加载
# asr_model = whisper.load_model("large")
# 使用字典缓存不同模型
asr_models = {}

def get_asr_model(model_size):
    """
    根据 model_size 加载 Whisper 模型，缓存已加载的模型
    """
    global asr_models
    if model_size not in asr_models:
        asr_models[model_size] = whisper.load_model(model_size)
    return asr_models[model_size]

# --------------------------------------------------------------------------
# 辅助函数：格式化文件大小
def format_file_size(file_path):
    size_bytes = os.path.getsize(file_path)
    if size_bytes < 1024 * 1024:
        kb = size_bytes / 1024
        return f"{kb:.2f} KB"
    else:
        mb = size_bytes / (1024 * 1024)
        return f"{mb:.2f} MB"

# --------------------------------------------------------------------------
# 核心功能函数
def download_models():
    model_list = [
        ("shape_predictor_68_face_landmarks.dat",
         "https://github.com/OpenTalker/SadTalker/releases/download/v0.0.2/shape_predictor_68_face_landmarks.dat",
         "checkpoints"),
        ("wav2lip.pth", "https://huggingface.co/guoyww/facevid2vid/resolve/main/wav2lip.pth", "checkpoints"),
        ("mapping_00109-model.pth.tar",
         "https://huggingface.co/guoyww/facevid2vid/resolve/main/mapping_00109-model.pth.tar", "checkpoints"),
        ("parsing_model.pth", "https://huggingface.co/guoyww/facevid2vid/resolve/main/parsing_model.pth", "checkpoints"),
        ("GFPGANv1.4.pth", "https://github.com/TencentARC/GFPGAN/releases/download/v1.3.8/GFPGANv1.4.pth",
         "checkpoints/gfpgan")
    ]
    for name, url, folder in model_list:
        os.makedirs(folder, exist_ok=True)
        dest = os.path.join(folder, name)
        if os.path.exists(dest):
            print(f"[已存在] {name}")
            continue
        print(f"[下载中] {name} ...")
        with requests.get(url, stream=True) as r:
            r.raise_for_status()
            with open(dest, 'wb') as f:
                for chunk in r.iter_content(chunk_size=8192):
                    f.write(chunk)
        print(f"[完成] {dest}")

def move_file_to_uploads(original_path, file_type="unknown"):
    if not original_path or not os.path.exists(original_path):
        return original_path

    abs_uploads = os.path.abspath(UPLOADS_DIR)
    abs_original = os.path.abspath(original_path)
    # 如果文件已在 UPLOADS_DIR 内，则直接返回
    if abs_original.startswith(abs_uploads):
        return original_path

    base_name = os.path.basename(original_path)
    new_name = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{file_type}_{base_name}"
    new_path = os.path.join(UPLOADS_DIR, new_name)
    try:
        shutil.copy2(original_path, new_path)
        size_kb = os.path.getsize(new_path) / 1024
        app_logger.info(f"Uploaded {file_type} -> {new_path} ({size_kb:.2f} KB)")
        return new_path
    except Exception as e:
        app_logger.error(f"move_file_to_uploads error: {e}")
        return original_path

def generate_speech(text):
    output_path = "output.wav"
    tts.tts_to_file(text=text, file_path=output_path)
    trim_tail_by_energy_and_gradient(output_path)
    return output_path

def transcribe_audio(audio_file, model_size):
    if not audio_file:
        return "⚠️ 没有上传任何音频文件"
    try:
        if not os.path.exists(audio_file):
            return "❗ 找不到音频文件，请重新上传"
        if os.path.getsize(audio_file) < 2048:
            return "⚠️ 音频文件太小，可能上传不完整或为空，请重新上传"
        import soundfile as sf
        try:
            _ = sf.info(audio_file)
        except Exception:
            return "⚠️ 音频文件格式不支持或内容损坏，请重新上传"

        # 将文件复制到 /uploads/（若之前已复制则直接使用）
        new_path = move_file_to_uploads(audio_file, file_type="audio")
        size_str = format_file_size(new_path)

        # 根据选定的模型尺寸加载模型
        model = get_asr_model(model_size)
        result = model.transcribe(new_path, language="zh")

        simplified = ""
        for char in result["text"]:
            if char in "。！？；，、,.!?;:":
                simplified += char
            else:
                simplified += cc.convert(char)

        # 记录识别日志，写明使用的模型和识别结果
        app_logger.info(f"语音识别使用模型: {model_size}，识别结果: {simplified}")
        asr_logger.info(f"语音识别使用模型: {model_size}，识别结果: {simplified}")  # 新增记录到 recognized.log
        save_recognition_history(result["text"], simplified, model_size)
        return f"识别结果（文件大小: {size_str}）：\n{simplified}"
    except Exception as e:
        raise e

def generate_video(image_path, audio_path):
    if not image_path or not os.path.exists(image_path):
        return "⚠️ 没有上传头像图片或文件不存在"
    if os.path.getsize(image_path) < 2048:
        return "⚠️ 上传的头像文件太小，可能无效"
    if not audio_path or not os.path.exists(audio_path):
        return "⚠️ 没有上传音频文件或文件不存在"
    if os.path.getsize(audio_path) < 2048:
        return "⚠️ 音频文件太小，可能无效或上传不完整"

    new_audio_path = move_file_to_uploads(audio_path, file_type="audio")

    output_dir = "results"
    os.makedirs(output_dir, exist_ok=True)
    launcher_path = os.path.abspath("sadtalker/launcher.py")
    cmd = [
        "python", launcher_path,
        "--driven_audio", new_audio_path,
        "--source_image", image_path,
        "--result_dir", output_dir,
        "--preprocess", "full",
        "--still",
        "--enhancer", "gfpgan"
    ]
    try:
        subprocess.run(cmd, check=True)
    except subprocess.CalledProcessError as e:
        return f"生成失败：\n命令：{' '.join(e.cmd)}\n返回码：{e.returncode}"

    output_video_path = os.path.join(output_dir, "result.mp4")
    return output_video_path if os.path.exists(output_video_path) else "生成失败，未找到视频文件"

def save_recognition_history(text_raw, text_simplified, model_used):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename_txt = os.path.join(RECOGNIZED_DIR, f"recognized_{timestamp}_{model_used}.txt")
    filename_docx = os.path.join(RECOGNIZED_DIR, f"recognized_{timestamp}_{model_used}.docx")

    with open(filename_txt, "w", encoding="utf-8") as f:
        f.write(f"[识别时间] {timestamp}\n")
        f.write(f"[使用模型] {model_used}\n")
        f.write(f"[原始文本]\n{text_raw}\n\n")
        f.write(f"[简体结果]\n{text_simplified}\n")

    from docx import Document
    doc = Document()
    doc.add_heading("语音识别结果", level=1)
    doc.add_paragraph(f"识别时间: {timestamp}")
    doc.add_paragraph(f"使用模型: {model_used}")
    doc.add_paragraph("原始文本:")
    doc.add_paragraph(text_raw)
    doc.add_paragraph("转换为简体：")
    doc.add_paragraph(text_simplified)
    doc.save(filename_docx)

def export_recognition_zip():
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    zip_path = os.path.join(RECOGNIZED_EXPORT_DIR, f"recognized_export_{timestamp}.zip")
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for filename in os.listdir(RECOGNIZED_DIR):
            file_path = os.path.join(RECOGNIZED_DIR, filename)
            zipf.write(file_path, arcname=filename)
    return zip_path

def search_history_by_question(query):
    hits = []
    for filename in os.listdir(RECOGNIZED_DIR):
        path = os.path.join(RECOGNIZED_DIR, filename)
        if filename.endswith(".txt"):
            with open(path, "r", encoding="utf-8") as f:
                content = f.read()
                if query in content:
                    hits.append(f"📄 {filename}\n{content[:300]}...\n---")
    if not hits:
        return "未找到相关内容。请尝试输入更常见的关键词。"
    return "\n\n".join(hits)

demo = gr.Blocks(css=material_css)

with demo:
    gr.Markdown("""
    <h2 style="text-align:center; color:#0abab5; font-weight:bold; margin-bottom:0.5em;">
    🎤 吉安智能体
    </h2>
    """)

    with gr.Tab("文字转语音"):
        text_input = gr.Textbox(label="输入文字")
        generate_btn = gr.Button("🎧 合成语音")
        output_audio = gr.Audio(label="语音文件", type="filepath", interactive=True)
        generate_btn.click(fn=generate_speech, inputs=text_input, outputs=output_audio)

    with gr.Tab("语音转文字"):
        with gr.Row():
            audio_input = gr.File(label="上传语音 (仅限 WAV 格式)", file_types=[".wav"], interactive=True)
            upload_status = gr.Textbox(label="语音上传状态", interactive=False, max_lines=1,
                                       container=True, show_copy_button=True)
        # 下拉菜单支持更多模型选项
        model_selector = gr.Dropdown(choices=["tiny", "base", "small", "medium", "large"],
                                     value="large", label="选择识别模型")
        transcribe_btn = gr.Button("📑 识别")
        asr_output = gr.Textbox(label="识别结果")

        def check_audio_upload_status(audio_file):
            if not audio_file:
                return ""
            if isinstance(audio_file, str) and os.path.exists(audio_file) and audio_file.endswith('.wav'):
                if os.path.getsize(audio_file) >= 2048:
                    size_str = format_file_size(audio_file)
                    return f"✅ 音频上传完成 (大小: {size_str})"
                else:
                    return "⚠️ 音频文件太小，可能无效"
            return "⚠️ 请上传 WAV 格式且大于2KB 的音频文件"

        audio_input.change(fn=check_audio_upload_status, inputs=audio_input, outputs=upload_status)
        # 识别按钮同时传入音频文件和模型选择
        transcribe_btn.click(fn=transcribe_audio, inputs=[audio_input, model_selector], outputs=asr_output)

    with gr.Tab("数字人动画"):
        with gr.Row():
            with gr.Column():
                image_input = gr.File(label="上传头像 (PNG/JPG)", file_types=[".png", ".jpg", ".jpeg"],
                                      interactive=True)
                image_name = gr.Textbox(label="头像文件名", interactive=False, max_lines=1)
                image_status = gr.Textbox(label="头像上传状态", interactive=False, max_lines=1,
                                          container=True, show_copy_button=True)
                image_preview = gr.Image(label="头像预览", interactive=False)

                def update_image_preview(image_file):
                    if not image_file or not os.path.exists(image_file):
                        return gr.update(visible=False, label="")
                    if os.path.getsize(image_file) < 2048 or not image_file.lower().endswith((".png", ".jpg", ".jpeg")):
                        return gr.update(visible=False, label="")
                    new_path = move_file_to_uploads(image_file, file_type="image")
                    im = PILImage.open(new_path)
                    w, h = im.size
                    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    return gr.update(
                        value=new_path,
                        visible=True,
                        label=f"分辨率: {w}x{h}  |  上传时间: {ts}"
                    )

                image_input.change(fn=update_image_preview, inputs=image_input, outputs=image_preview)

        with gr.Row():
            with gr.Column():
                driven_audio_input = gr.Audio(label="使用合成或自己语音", type="filepath", interactive=True)
                audio_name = gr.Textbox(label="音频文件名", interactive=False, max_lines=1)
                audio_status = gr.Textbox(label="音频上传状态", interactive=False, max_lines=1,
                                          container=True, show_copy_button=True)

        generate_video_btn = gr.Button("🎥 生成动画")
        video_output = gr.Video(label="数字人视频")

        def check_image_upload_status(image_file):
            if not image_file:
                return ""
            if isinstance(image_file, str) and os.path.exists(image_file):
                if os.path.getsize(image_file) >= 2048 and image_file.lower().endswith(('.png', '.jpg', '.jpeg')):
                    size_str = format_file_size(image_file)
                    return f"✅ 头像上传完成 (大小: {size_str})"
                else:
                    return "⚠️ 头像文件太小或格式不正确"
            return ""

        def check_audio_upload_status_generic(audio_file):
            if not audio_file:
                return ""
            if isinstance(audio_file, str) and os.path.exists(audio_file):
                if os.path.getsize(audio_file) >= 2048:
                    size_str = format_file_size(audio_file)
                    return f"✅ 音频上传完成 (大小: {size_str})"
                else:
                    return "⚠️ 音频文件太小或格式不正确"
            return ""

        image_input.change(fn=lambda f: os.path.basename(f) if f else "未选择文件", inputs=image_input,
                           outputs=image_name)
        image_input.change(fn=check_image_upload_status, inputs=image_input, outputs=image_status)
        driven_audio_input.change(fn=lambda f: os.path.basename(f) if f else "未选择文件", inputs=driven_audio_input,
                                  outputs=audio_name)
        driven_audio_input.change(fn=check_audio_upload_status_generic, inputs=driven_audio_input, outputs=audio_status)

        generate_video_btn.click(fn=generate_video, inputs=[image_input, driven_audio_input], outputs=video_output)

    with gr.Tab("模型下载"):
        gr.Markdown("### 🧩 首次使用请点击下载 SadTalker 模型")
        download_btn = gr.Button("📥 下载模型")
        download_output = gr.Textbox(label="状态输出")
        download_btn.click(fn=download_models, outputs=download_output)

    with gr.Tab("识别历史"):
        gr.Markdown("### 📄 导出历史 / 查询内容")
        with gr.Row():
            export_btn = gr.Button("📦 导出 ZIP")
            export_file = gr.File(label="下载识别记录压缩包")
            export_btn.click(fn=export_recognition_zip, outputs=export_file)
        with gr.Row():
            query_input = gr.Textbox(label="输入关键词或内容问题")
            query_btn = gr.Button("🔍 查询记录")
            query_result = gr.Textbox(label="查询结果", lines=8)
            query_btn.click(fn=search_history_by_question, inputs=query_input, outputs=query_result)

if __name__ == "__main__":
    demo.launch()
